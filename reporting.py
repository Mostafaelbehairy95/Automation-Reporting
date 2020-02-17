from docx import *
import openpyxl

# Position of item inside word table
posWordReport = [
    [2, 1, 1, 2],  # Request No.
    [4, 1, 2, 2],  # Tester
    [3, 1, 5, 3],  # Test Date
    [4, 3, 5, 1],  # Test Case No.
    [5, 1, 5, 2],  # Test Case Description
    [6, 1, 5, 7],  # Results
    [6, 3, 5, 8],  # Incident No
    [8, 1, 5, 4],  # Requirement(s)
    [9, 1, 5, 9],  # Roles & Responsibilities
    [10, 1, 5, 10],  # Set up procedures
    [11, 1, 5, 11],  # Stop Procedures
    [13, 1, 5, 12],  # Hardware
    [14, 1, 5, 13],  # Software
    [15, 1, 5, 14],  # Procedural Requirements
    [17, 1, 5, 15],  # Test Items and Features:
    [18, 1, 5, 16],  # Input Specifications:
    [19, 1, 5, 17],  # Procedural steps:
    [20, 1, 5, 5],  # Expected Results of Case:
    [22, 1, 5, 18]  # Output Specifications
]
posWordIncident = [
    [2, 1, 1, 2],   # Request No.
    [4, 1, 2, 2],   # Tester
    [3, 1, 5, 3],   # Test Date
    [4, 3, 5, 1],   # Test Case No.
    [5, 1, 5, 2],   # Test Case Description
    [6, 1, 5, 8],   # Incident No
    [8, 1, 5, 13],  # Environmental Information:
    [9, 1, 5, 19],  # Unusual Events
    [11, 1, 5, 20], # Summary of Incident
    [12, 1, 5, 16], # input
    [13, 1, 5, 5],  # Expected Results:
    [15, 1, 5, 21], # Actual Results
    [16, 1, 5, 22], # Abnormalities
    [17, 1, 5, 17] # Procedural Step:
]

wordPath = input("Enter path of Word: ")
excelPath = input("Enter path of Excel: ")
tab = input("Enter tabName which contain log: ")
reportName = input("write Name of word which you want to save: ")
testCase = input("Enter Number of Test Case: ")

# Configuration
word = Document(wordPath)
excel = openpyxl.load_workbook(excelPath)
excel_sh = excel[tab]
print("ـــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــــ")
print("Tab Name: ",excel_sh.title)


# Code
print("(=", end="")
for x in range(0,int(testCase)):
    word.tables[x].cell(posWordReport[0][0], posWordReport[0][1]).text = str(
        excel_sh.cell(posWordReport[0][2], posWordReport[0][3]).value)

    word.tables[x].cell(posWordReport[1][0], posWordReport[1][1]).text = str(
        excel_sh.cell(posWordReport[1][2], posWordReport[1][3]).value)
    for i in range(2,19):
        word.tables[x].cell(posWordReport[i][0],posWordReport[i][1]).text = str(
            excel_sh.cell(posWordReport[i][2] + x,posWordReport[i][3]).value)
    print("=", end="")

print("=)", end = "Done")

word.save(reportName)


#
#